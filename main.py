from docx import Document
import os


def term_cel():
    marca_cell = input("Marca: ")
    modelo_cell = input("Modelo: ")
    imei_cell = input("IMEI: ")
    ns_cell = input("Numero de Serie: ")
    estado_cell = input("Estado de conservação equipamento: ")
    linha_cell = input("Chip telefonia Movel: ")
    
    dados_especificos = {
        "<<marca>>" : marca_cell,
         "<<modelo>>" : modelo_cell,
         "<<imei>>" : imei_cell,
         "<<ns>>" : ns_cell,
         "<<estado>>" : estado_cell,
         "<<chip>>" : linha_cell,
         "<<hostname>>" : ""
    }
    return dados_especificos
 
def term_computer():
    modelo_computer = input("Modelo: ")
    ns_computer = input("Numero de serie: ")
    host_name = input("Hostname: ")
    estado_computer = input("Estado do equipamento: ")
    
    dados_especificos = {
        "<<modelo>>" : modelo_computer,
         "<<ns>>" : ns_computer,
         "<<estado>>" : estado_computer,
         "<<hostname>>" : host_name
    }
    return dados_especificos

def term_perifericos():
    nome_equip = input("Nome do equipamento: ")
    marca_equip = input("Marca do equipamento: ")
    estado_equip = input("Estado do equipamento: ")
    
    dados_especificos = {
        "<<nome>>" : nome_equip,
        "<<marca>>" : marca_equip,
        "<<estado>>" : estado_equip
    }
    return dados_especificos

       
while True:
    print("\n---- Menu ----")
    print("Escolha modelo de termo: ")
    print("1. celular ")
    print("2. Computador ")
    print("3. perifericos ")
    print("4. Sair ")
    
    escolha = input("digite uma opção: ")
    
    if escolha in ["1", "2", "3"]:
        print("Dados do colaborador")
        nome_colab = input("nome: ")
        user_ad = input("Usuario AD: ")
        
        dados_comuns = {
            "<<nome>>" : nome_colab,
            "<<user>>" : user_ad
        }
        
        nome_template = ""
        dados_especificos = {}
        
        if escolha ==  "1":
            nome_template = "termo_teste_cell.docx"
            dados_especificos = term_cel()
        elif escolha == "2":
            nome_template = "termo_teste_note.docx"
            dados_especificos = term_computer()
        elif escolha == "3":
            nome_template = "termo_teste_perifericos.docx"
            dados_especificos = term_perifericos()
        
        dados_finais = dados_comuns.copy()
        dados_finais.update(dados_especificos)
        
        print("\n Validando dados finais...")
        print(dados_finais)
        
        print("Gerando termo...")
        
        try:
            #carregar o doc
            doc =  Document(nome_template)
            
            #transforma os dados finais em chaves e valorres, dps percorre o documento
            for tag, valor in dados_finais.items():
                #substitui nos paragrafos
                for p in doc.paragraphs:
                    p.text = p.text.replace(tag, str(valor))
                    #substitui nas tabelas
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell.text = cell.text.replace(tag, str(valor))
                            
            nome_arquivo = f"{nome_colab.replace(' ', ' ')}"
            output_folder = "Termo_Gerado"
            
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
                
            path_docx = os.path.join(output_folder, f"{nome_arquivo}.docx")
            path_pdf = os.path.join(output_folder, f"{nome_arquivo}.pdf")
            
            doc.save(path_docx)
            print(f"Documento salvo em Word salvo em: {path_docx}")
            
        except Exception as e:
            print(f"Erro ao gerar termo: {e}")
            print(f"Verifique se o arquivo está na pasta correta")
            
    elif escolha == "4":
        print("saindo do programa")
        break
    else:
        print("opção invalida")
